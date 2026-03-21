"""Convert TFP_KillTracker_Technical_Documentation.docx → .pdf using ReportLab."""

import docx
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, KeepTogether, HRFlowable
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# ── Paths ─────────────────────────────────────────────────────────────────────
DOCX_PATH = '/home/user/TFP_KillTracker/TFP_KillTracker_Technical_Documentation.docx'
PDF_PATH  = '/home/user/TFP_KillTracker/TFP_KillTracker_Technical_Documentation.pdf'

# ── Colour palette (matches the docx) ────────────────────────────────────────
DARK    = colors.HexColor('#2E4057')
MID     = colors.HexColor('#555555')
HDR_BG  = colors.HexColor('#2E4057')
ALT_BG  = colors.HexColor('#F0F4FA')
CODE_BG = colors.HexColor('#F3F3F3')
WHITE   = colors.white
BLACK   = colors.black

# ── Styles ────────────────────────────────────────────────────────────────────
base = ParagraphStyle('base', fontName='Helvetica', fontSize=9,
                      leading=13, spaceAfter=4)

styles = {
    'title': ParagraphStyle('title', parent=base,
                             fontName='Helvetica-Bold', fontSize=26,
                             textColor=DARK, alignment=TA_CENTER,
                             spaceAfter=10),
    'subtitle': ParagraphStyle('subtitle', parent=base,
                                fontName='Helvetica', fontSize=14,
                                textColor=MID, alignment=TA_CENTER,
                                spaceAfter=8),
    'version': ParagraphStyle('version', parent=base,
                               fontName='Helvetica', fontSize=10,
                               textColor=MID, alignment=TA_CENTER,
                               spaceAfter=6),
    'h1': ParagraphStyle('h1', parent=base,
                          fontName='Helvetica-Bold', fontSize=15,
                          textColor=DARK, spaceBefore=14, spaceAfter=6,
                          borderPad=2),
    'h2': ParagraphStyle('h2', parent=base,
                          fontName='Helvetica-Bold', fontSize=11,
                          textColor=DARK, spaceBefore=10, spaceAfter=4),
    'normal': ParagraphStyle('normal', parent=base,
                              fontName='Helvetica', fontSize=9,
                              leading=13, spaceAfter=5),
    'bullet': ParagraphStyle('bullet', parent=base,
                              fontName='Helvetica', fontSize=9,
                              leading=13, leftIndent=14, spaceAfter=3,
                              bulletIndent=4, bulletText='•'),
    'numbered': ParagraphStyle('numbered', parent=base,
                                fontName='Helvetica', fontSize=9,
                                leading=13, leftIndent=18, spaceAfter=3),
    'code': ParagraphStyle('code', parent=base,
                            fontName='Courier', fontSize=8,
                            leading=11, leftIndent=18, spaceAfter=1,
                            backColor=CODE_BG),
}


def para_style(style_name: str):
    """Map docx style name → our ParagraphStyle."""
    if style_name == 'Normal':          return styles['normal']
    if style_name == 'Heading 1':       return styles['h1']
    if style_name == 'Heading 2':       return styles['h2']
    if style_name == 'List Bullet':     return styles['bullet']
    if style_name == 'List Number':     return styles['numbered']
    if style_name == 'No Spacing':      return styles['code']
    return styles['normal']


def is_title_block(para, doc):
    """True for the first three non-empty centred paragraphs (title page)."""
    idx = list(doc.paragraphs).index(para)
    return idx < 10 and para.alignment == 1  # WD_ALIGN_PARAGRAPH.CENTER == 1


def safe(txt: str) -> str:
    """Escape characters that would break ReportLab's XML parser."""
    return (txt.replace('&', '&amp;')
               .replace('<', '&lt;')
               .replace('>', '&gt;')
               .replace('"', '&quot;'))


def make_table(tbl):
    """Convert a docx Table → ReportLab Table flowable."""
    data = []
    for ri, row in enumerate(tbl.rows):
        r = []
        for cell in row.cells:
            txt = safe(cell.text.strip())
            if ri == 0:
                r.append(Paragraph(txt, ParagraphStyle(
                    'thdr', fontName='Helvetica-Bold', fontSize=8,
                    textColor=WHITE, leading=11)))
            else:
                r.append(Paragraph(txt, ParagraphStyle(
                    'tcell', fontName='Helvetica', fontSize=8,
                    textColor=BLACK, leading=11)))
        data.append(r)

    if not data:
        return None

    ncols = len(data[0])
    col_w = (A4[0] - 3*cm) / ncols  # distribute evenly within margins

    ts = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), HDR_BG),
        ('TEXTCOLOR',  (0, 0), (-1, 0), WHITE),
        ('FONTNAME',   (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE',   (0, 0), (-1, -1), 8),
        ('ALIGN',      (0, 0), (-1, 0), 'CENTER'),
        ('VALIGN',     (0, 0), (-1, -1), 'TOP'),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [WHITE, ALT_BG]),
        ('GRID',       (0, 0), (-1, -1), 0.4, colors.HexColor('#CCCCCC')),
        ('LEFTPADDING',  (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('TOPPADDING',   (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING',(0, 0), (-1, -1), 3),
    ])

    t = Table(data, colWidths=[col_w] * ncols, repeatRows=1)
    t.setStyle(ts)
    return t


# ── Read docx ─────────────────────────────────────────────────────────────────
doc = docx.Document(DOCX_PATH)

# Build a list of (type, object): ('para', para) or ('table', table)
elements_src = []
table_set = {id(t): t for t in doc.tables}
table_done = set()

for block in doc.element.body:
    tag = block.tag.split('}')[-1]
    if tag == 'p':
        # find matching paragraph
        for p in doc.paragraphs:
            if p._p is block:
                elements_src.append(('para', p))
                break
    elif tag == 'tbl':
        for t in doc.tables:
            if t._tbl is block and id(t) not in table_done:
                elements_src.append(('table', t))
                table_done.add(id(t))
                break

# ── Build PDF flowables ───────────────────────────────────────────────────────
story = []
title_paras_seen = 0

numbered_counters = {}   # track list number per indent level

for kind, obj in elements_src:
    if kind == 'table':
        tbl = make_table(obj)
        if tbl:
            story.append(Spacer(1, 4))
            story.append(tbl)
            story.append(Spacer(1, 6))
        numbered_counters.clear()
        continue

    # ── Paragraph ──
    p = obj
    text = safe(p.text.strip())
    sname = p.style.name

    # Page breaks
    for run in p.runs:
        if run._r.xml.find('w:br') != -1 or '<w:br' in run._r.xml:
            story.append(PageBreak())
            break
    else:
        # Check paragraph-level page break
        pPr = p._p.find(qn('w:pPr'))
        if pPr is not None:
            pgBr = pPr.find(qn('w:pageBreakBefore'))
            if pgBr is not None and pgBr.get(qn('w:val'), 'true') != 'false':
                story.append(PageBreak())

    if not text:
        story.append(Spacer(1, 5))
        continue

    # Title page detection (first few centred non-heading paragraphs)
    if sname == 'Normal' and p.alignment == 1 and title_paras_seen < 3:
        if title_paras_seen == 0:
            s = styles['title']
        elif title_paras_seen == 1:
            s = styles['subtitle']
        else:
            s = styles['version']
        title_paras_seen += 1
        story.append(Spacer(1, 8))
        story.append(Paragraph(text, s))
        continue

    s = para_style(sname)

    if sname == 'Heading 1':
        story.append(Spacer(1, 8))
        story.append(HRFlowable(width='100%', thickness=1, color=DARK, spaceAfter=4))
        story.append(Paragraph(text, s))
        numbered_counters.clear()
    elif sname == 'Heading 2':
        story.append(Paragraph(text, s))
        numbered_counters.clear()
    elif sname == 'List Number':
        lvl = 0
        numbered_counters[lvl] = numbered_counters.get(lvl, 0) + 1
        n = numbered_counters[lvl]
        story.append(Paragraph(f'{n}.&nbsp;&nbsp;{text}', s))
    elif sname == 'List Bullet':
        story.append(Paragraph(text, s))
    elif sname == 'No Spacing':
        story.append(Paragraph(text.replace(' ', '&nbsp;'), styles['code']))
    else:
        story.append(Paragraph(text, s))

# ── Render ────────────────────────────────────────────────────────────────────
pdf = SimpleDocTemplate(
    PDF_PATH,
    pagesize=A4,
    leftMargin=2.5*cm, rightMargin=2.5*cm,
    topMargin=2.0*cm, bottomMargin=2.0*cm,
    title='TFP Kill Tracker — Technical Documentation',
    author='TFP Kill Tracker',
)
pdf.build(story)
print(f'Saved: {PDF_PATH}')
