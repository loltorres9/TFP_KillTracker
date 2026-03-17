"""Generate TFP Kill Tracker technical documentation as a .docx file."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h

def add_para(doc, text='', bold=False, italic=False, size=None, color=None, space_before=None, space_after=None):
    p = doc.add_paragraph()
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def add_code(doc, text):
    p = doc.add_paragraph(style='No Spacing')
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    # light gray shading on paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F3F3F3')
    pPr.append(shd)
    return p

def add_table(doc, headers, rows, header_bg='2E4057', header_fg='FFFFFF', alt_bg='F0F4FA'):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(*bytes.fromhex(header_fg))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg = alt_bg if ri % 2 == 1 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
    return tbl

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def numbered(doc, text, level=0):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

# ── Document ──────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Styles ──
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10)

# ═══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('TFP KILL TRACKER')
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

doc.add_paragraph()

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run('Full Technical Documentation')
rs.font.size = Pt(16)
rs.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
ver.add_run(f'Version 1.0  ·  {datetime.date.today().strftime("%d %B %Y")}').font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  1. OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '1. Overview', 1)

add_para(doc,
    'TFP Kill Tracker is a fully client-side, single-page web application that fetches '
    'live kill statistics from a public Google Sheets CSV export and renders them as a '
    'sortable, filterable leaderboard. There is no back-end server; all processing — '
    'CSV parsing, aggregation, filtering, and DOM rendering — occurs in the browser.',
    space_after=6)

add_heading(doc, '1.1 Technology Stack', 2)
add_table(doc,
    ['Layer', 'Technology', 'Notes'],
    [
        ['Front-end', 'Vanilla HTML / CSS / JavaScript (ES6+)', 'No frameworks or build tools'],
        ['Fonts', 'Google Fonts – Oswald & Open Sans', 'Loaded via CDN'],
        ['Data Source', 'Google Sheets public CSV export', 'Fetched at page load via fetch()'],
        ['Deployment', 'GitHub Actions → GitHub Pages', 'Automatic on push to main / claude/*'],
        ['Documentation', 'python-docx', 'This document'],
    ]
)

doc.add_paragraph()
add_heading(doc, '1.2 Key Design Principles', 2)
bullet(doc, 'Zero dependencies — no npm, no bundler, no framework.')
bullet(doc, 'Single file — the entire application lives in index.html (CSS + JS inlined).')
bullet(doc, 'Stateless server — all state is in-memory JavaScript variables; nothing is persisted.')
bullet(doc, 'Re-aggregate on every filter change — guarantees all derived metrics stay consistent.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  2. ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '2. Architecture', 1)

add_heading(doc, '2.1 File Structure', 2)
add_table(doc,
    ['File', 'Purpose'],
    [
        ['index.html', 'Entire application — HTML structure, inline CSS, inline JavaScript'],
        ['favicon.svg', 'Browser tab icon (SVG crosshair graphic)'],
        ['README.md', 'User-facing feature overview and deployment guide'],
        ['.github/workflows/', 'GitHub Actions CI/CD pipeline for GitHub Pages deployment'],
    ]
)

doc.add_paragraph()
add_heading(doc, '2.2 Data Flow', 2)
add_para(doc, 'The application follows a strict linear pipeline on every filter change:', space_after=4)
numbered(doc, 'FETCH — GET request to Google Sheets CSV export URL (CSV_URL constant).')
numbered(doc, 'PARSE — Raw text split into lines; each line processed by parseCSVLine() to handle quoted fields.')
numbered(doc, 'AGGREGATE (full) — buildAggregates() loops all rows once and builds aggPlayers, a keyed object of career totals per player.')
numbered(doc, 'FILTER — applyFilters() re-loops rawRows applying four independent filters (event type, player selection, mission selection, Zeus status), building a temporary filteredPlayers array.')
numbered(doc, 'RENDER — Five render functions consume filteredPlayers and write HTML to specific DOM targets.')
doc.add_paragraph()
add_para(doc,
    'Steps 4–5 re-execute on every user interaction (filter toggle, pill click, column sort header click). '
    'Steps 1–3 run once at page load.',
    italic=True, space_after=6)

add_heading(doc, '2.3 Global State Variables', 2)
add_table(doc,
    ['Variable', 'Type', 'Default', 'Purpose'],
    [
        ['rawRows', 'Array<Object>', '[]', 'Every CSV row parsed into {header: value} objects'],
        ['aggPlayers', 'Object', '{}', 'Full career aggregates keyed by player name'],
        ['filteredPlayers', 'Array<Object>', '[]', 'Aggregated stats for the current filter set'],
        ['selectedPlayers', 'Set | null', 'null', 'null = all players; Set = explicit selection'],
        ['selectedMissions', 'Set | null', 'null', 'null = all missions; Set = explicit selection'],
        ['showJointOps', 'boolean', 'true', 'Include Joint Op events'],
        ['showRegularEvents', 'boolean', 'true', 'Include regular op events'],
        ['zeusFilter', 'string', '"all"', '"all" | "no-zeus" | "zeus-only"'],
        ['infSortCol', 'number', '2', 'Active sort column index for infantry table'],
        ['infSortAsc', 'boolean', 'false', 'Sort direction for infantry table'],
        ['vehSortCol', 'number', '2', 'Active sort column index for vehicle table'],
        ['vehSortAsc', 'boolean', 'false', 'Sort direction for vehicle table'],
        ['currentModalPlayer', 'string', '""', 'Player name currently shown in detail modal'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  3. DATA SOURCE
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '3. Data Source', 1)

add_heading(doc, '3.1 CSV Configuration', 2)
add_para(doc,
    'The application fetches from the URL defined by the constant CSV_URL near the top of index.html. '
    'To point the tracker at a different sheet, update this constant. The sheet must be published '
    'publicly ("File → Share → Publish to web → CSV").',
    space_after=6)

add_heading(doc, '3.2 Expected Column Schema', 2)
add_para(doc, 'One row per player per mission. Columns must appear in this order:', space_after=4)
add_table(doc,
    ['Index', 'Column Name', 'Description'],
    [
        ['0', 'Source File', 'Filename encoding date (YYYY_MM_DD__...) and mission name'],
        ['1', 'Mission', 'Human-readable mission name, expected to contain "(YYYY-MM-DD)"'],
        ['2', 'World', 'Map/terrain name'],
        ['3', 'Username', 'Player name (used as aggregate key)'],
        ['4', 'Side', 'Faction (e.g. BLUFOR, OPFOR)'],
        ['5', 'Group', 'Role/squad; "zeus" group triggers Zeus classification'],
        ['6', 'Kills (On Foot)', 'Infantry kills'],
        ['7', 'Deaths (On Foot)', 'Infantry deaths'],
        ['8', 'K/D (On Foot)', 'Ignored — recomputed from kills/deaths'],
        ['9', 'Teamkills (On Foot)', 'Friendly fire kills (infantry)'],
        ['10', 'Shots (On Foot)', 'Rounds fired on foot'],
        ['11', 'Hits Taken (On Foot)', 'Times hit while on foot'],
        ['12', 'Shots/Kill (On Foot)', 'Ignored — recomputed'],
        ['13', 'Avg Kill Dist On Foot (m)', 'Average kill distance (on foot)'],
        ['14', 'Longest Kill On Foot (m)', 'Longest single kill distance (on foot)'],
        ['15', 'Kills (In Vehicle)', 'Vehicle kills'],
        ['16', 'Deaths (In Vehicle)', 'Deaths while in vehicle'],
        ['17', 'K/D (In Vehicle)', 'Ignored — recomputed'],
        ['18', 'Teamkills (In Vehicle)', 'Friendly fire from vehicle'],
        ['19', 'Vehicle Kills (On Foot)', 'Vehicles destroyed while on foot'],
        ['20', 'Vehicle Kills (In Vehicle)', 'Vehicles destroyed while in vehicle'],
        ['21', 'Shots (In Vehicle)', 'Rounds fired from vehicle'],
        ['22', 'Hits Taken (In Vehicle)', 'Times vehicle was hit'],
        ['23', 'Shots/Kill (In Vehicle)', 'Ignored — recomputed'],
        ['24', 'Avg Kill Dist In Vehicle (m)', 'Average kill distance (vehicle)'],
        ['25', 'Longest Kill In Vehicle (m)', 'Longest single kill distance (vehicle)'],
        ['26', 'Weapon Kills (JSON)', 'JSON object: {"WeaponName": killCount, ...}'],
        ['27', 'Top Weapon', 'Weapon used for longest kill (used in award card)'],
    ]
)

doc.add_paragraph()
add_heading(doc, '3.3 CSV Parsing', 2)
add_para(doc,
    'parseCSVLine(line) handles RFC-4180 compliant CSV including quoted fields that contain commas '
    'or embedded double-quotes (escaped as ""). The parser is character-by-character and maintains '
    'a boolean inQuote state. Quotes within quoted fields are unescaped before storage.',
    space_after=6)

add_heading(doc, '3.4 Mission Date Extraction', 2)
add_para(doc,
    'Mission names are expected to contain an ISO date in parentheses, e.g. "Operation Nightfall (2024-03-15)". '
    'The missionDate(missionName) function applies the regex /\\((\\d{4}-\\d{2}-\\d{2})\\)/ '
    'and reformats the result to DD/MM/YYYY for display. If no date is found the original string is returned unchanged.',
    space_after=4)
add_code(doc, 'missionDate("Op Nightfall (2024-03-15)")  →  "15/03/2024"')

doc.add_paragraph()
add_heading(doc, '3.5 Joint Op Detection', 2)
add_para(doc,
    'isJointOp(filename) determines whether a mission should be classified as a Joint Op. '
    'The logic parses the YYYY_MM_DD prefix from the filename, determines the day-of-week, '
    'then checks whether that Saturday is the last Saturday of its calendar month, '
    'or whether the date is a Sunday immediately following the last Saturday.',
    space_after=4)
add_code(doc, '// filename format: "2024_03_30__MissionName.sqm"')
add_code(doc, '// Returns true if date is last Saturday or last-Saturday+1 (Sunday) of the month')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  4. AGGREGATION ENGINE
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '4. Aggregation Engine', 1)

add_heading(doc, '4.1 buildAggregates()', 2)
add_para(doc,
    'Called once after the CSV is fetched. Loops rawRows and builds aggPlayers — a dictionary '
    'keyed by player name. Each entry is initialised with zero values on first encounter, '
    'then accumulated across every row.',
    space_after=4)
add_para(doc, 'Computed fields per player:', space_after=2)
bullet(doc, 'kdFoot / kdVeh — kills ÷ deaths (Infinity if deaths = 0, 0 if kills = 0)')
bullet(doc, 'spkFoot / spkVeh — shots ÷ kills (null if kills = 0)')
bullet(doc, 'avgDistFoot / avgDistVeh — weighted mean: (Σ avgDist × kills) ÷ Σ kills')
bullet(doc, 'missionCount — cardinality of the missions Set')
bullet(doc, 'topRole — normalizeRole() applied to Group; role with highest count wins')
bullet(doc, 'longestKillWeapon — weapon column from the row that contained the max distance')

add_heading(doc, '4.2 Player Object Schema', 2)
add_code(doc, '{')
add_code(doc, '  name: string,')
add_code(doc, '  missions: Set<string>,       // unique mission names')
add_code(doc, '  worlds:   Set<string>,       // unique map names')
add_code(doc, '')
add_code(doc, '  // ── Infantry ──')
add_code(doc, '  killsOnFoot:   number,')
add_code(doc, '  deathsOnFoot:  number,')
add_code(doc, '  tkOnFoot:      number,')
add_code(doc, '  shotsOnFoot:   number,')
add_code(doc, '  hitsOnFoot:    number,')
add_code(doc, '  suicides:      number,')
add_code(doc, '  maxLongestFoot: number,      // metres')
add_code(doc, '  avgDistFoot:   number,       // weighted mean (metres)')
add_code(doc, '')
add_code(doc, '  // ── Vehicle ──')
add_code(doc, '  killsInVeh:    number,')
add_code(doc, '  deathsInVeh:   number,')
add_code(doc, '  tkInVeh:       number,')
add_code(doc, '  vehKillsFoot:  number,       // vehicles destroyed while on foot')
add_code(doc, '  vehKillsVeh:   number,       // vehicles destroyed while in vehicle')
add_code(doc, '  shotsInVeh:    number,')
add_code(doc, '  hitsInVeh:     number,')
add_code(doc, '  maxLongestVeh: number,')
add_code(doc, '  avgDistVeh:    number,')
add_code(doc, '')
add_code(doc, '  // ── Derived ──')
add_code(doc, '  kdFoot:        number,')
add_code(doc, '  kdVeh:         number,')
add_code(doc, '  spkFoot:       number | null,')
add_code(doc, '  spkVeh:        number | null,')
add_code(doc, '  missionCount:  number,')
add_code(doc, '')
add_code(doc, '  // ── Meta ──')
add_code(doc, '  weaponKills:        { [weapon: string]: number },')
add_code(doc, '  longestKillWeapon:  string,')
add_code(doc, '  roleCounts:         { [role: string]: number },')
add_code(doc, '  topRole:            string | null,')
add_code(doc, '  topRoleCount:       number,')
add_code(doc, '  missionRows:        Array<CSVRowObject>,')
add_code(doc, '}')

add_heading(doc, '4.3 Role Normalisation', 2)
add_para(doc,
    'normalizeRole(role) strips unit prefixes and cosmetic suffixes before role counting, '
    'so that "SL Alpha @player" and "SL Alpha" are counted as the same role.',
    space_after=4)
bullet(doc, 'Removes @username suffixes')
bullet(doc, 'Strips Alpha/Bravo/Charlie/Delta/Echo/Foxtrot squad letters')
bullet(doc, 'Removes numbered squad prefixes (S1, S2 …)')
bullet(doc, 'Removes colour names (Red, Blue, Green, Yellow …)')
bullet(doc, 'Removes NATO phonetic identifiers (Alpha, Bravo … Foxtrot)')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  5. FILTER SYSTEM
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '5. Filter System', 1)

add_heading(doc, '5.1 Filter Pipeline', 2)
add_para(doc,
    'applyFilters() is the central function that rebuilds filteredPlayers from rawRows. '
    'Filters are applied in sequence; a row must pass all active filters to contribute to '
    'the output aggregation.',
    space_after=4)

add_table(doc,
    ['Step', 'Filter', 'State Variable', 'Condition'],
    [
        ['1', 'Event Type', 'showJointOps / showRegularEvents', 'isJointOp(srcFile) must match the selected event type'],
        ['2', 'Player', 'selectedPlayers', 'null = all; else row.Username must be in the Set'],
        ['3', 'Mission', 'selectedMissions', 'null = all; else row.Mission must be in the Set'],
        ['4', 'Zeus', 'zeusFilter', 'Applied post-aggregation: checks if player has any Group="zeus" row'],
    ]
)

doc.add_paragraph()
add_heading(doc, '5.2 Pill Availability Cascade', 2)
add_para(doc,
    'Available pills are recomputed by refreshPills() after every filter change to prevent '
    'showing players or missions that have no data under the current filter set.',
    space_after=4)
numbered(doc, 'getEventFilteredRows() returns rows matching the event-type filter.')
numbered(doc, 'Available missions = unique Mission values in those rows.')
numbered(doc, 'Available players  = unique Username values in rows that also match the selected missions.')
numbered(doc, 'Both pill lists are re-rendered against the search input substring filter.')

add_heading(doc, '5.3 Player Pills vs Mission Pills', 2)
add_table(doc,
    ['Aspect', 'Player Pills', 'Mission Pills'],
    [
        ['Renderer', 'renderCareerPills()', 'renderSelectablePills()'],
        ['Left-click', 'Toggle player selection (filterToPlayer)', 'Toggle mission selection'],
        ['Right-click', 'Open full career page (openCareerPage)', 'No action'],
        ['Multi-select', 'Yes — multiple players can be active simultaneously', 'Yes — multiple missions can be active simultaneously'],
        ['Active state', 'pill.active class (dark background)', 'pill.active class (dark background)'],
        ['null state', 'null = all players active', 'null = all missions active'],
        ['Search', 'playerSearch input (case-insensitive substring)', 'missionSearch input (case-insensitive substring)'],
    ]
)

doc.add_paragraph()
add_heading(doc, '5.4 Zeus Filter', 2)
add_para(doc,
    'The Zeus filter is applied after re-aggregation. For each player in the filtered '
    'aggregation, the system checks whether any of their rawRows has Group (lowercased) === "zeus".',
    space_after=4)
add_table(doc,
    ['Value', 'Behaviour'],
    [
        ['"all"', 'Include all players regardless of Zeus status'],
        ['"no-zeus"', 'Exclude players who have at least one Zeus row in the filtered set'],
        ['"zeus-only"', 'Include only players who have at least one Zeus row'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  6. RENDERING FUNCTIONS
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '6. Rendering Functions', 1)

add_heading(doc, '6.1 renderStats()', 2)
add_para(doc, 'Renders the top stats bar with aggregate totals across all filteredPlayers.', space_after=4)
add_table(doc,
    ['Card Label', 'Value Source'],
    [
        ['Infantry Kills', 'Σ killsOnFoot'],
        ['Kills from Vehicles', 'Σ killsInVeh'],
        ['Total Deaths', 'Σ (deathsOnFoot + deathsInVeh)'],
        ['Teamkills', 'Σ (tkOnFoot + tkInVeh)'],
        ['Vehicles Destroyed', 'Σ (vehKillsFoot + vehKillsVeh)'],
        ['Players', 'filteredPlayers.length'],
        ['Missions', 'unique mission count across filtered rows'],
        ['Average K/D', '(Σ killsOnFoot) / (Σ deathsOnFoot), formatted to 2 d.p.'],
    ]
)

doc.add_paragraph()
add_heading(doc, '6.2 renderChart()', 2)
add_para(doc,
    'Renders a horizontal bar chart of the top 10 infantry killers. '
    'Bar widths are proportional to the highest kill count in the visible set. '
    'Each bar shows player name and kill count as a label.',
    space_after=6)

add_heading(doc, '6.3 renderLeader() — Award Cards', 2)
add_table(doc,
    ['Award', 'Icon', 'Metric', 'Eligibility', 'Display'],
    [
        ['Executioner', '👑', 'Most killsOnFoot', 'killsOnFoot > 0', 'Name · kills · K/D'],
        ['Pro Sniper', '🎯', 'Highest maxLongestFoot', 'Min 2 kills on foot', 'Name · distance (m) · avg dist · weapon'],
        ['Perfect Aim', '🔫', 'Lowest spkFoot (shots/kill)', 'Min 3 kills, spkFoot ≥ 1', 'Name · SPK · kill count'],
        ['K/D Player', '⚔️', 'Highest kdFoot', 'Min 3 kills on foot', 'Name · K/D · kills/deaths'],
    ]
)

doc.add_paragraph()
add_heading(doc, '6.4 renderInfantryTable()', 2)
add_para(doc, 'Renders the Infantry Combat Stats table. Columns:', space_after=4)
add_table(doc,
    ['#', 'Column', 'Key', 'Default Sort'],
    [
        ['0', 'Rank (#)', '_rank', '—'],
        ['1', 'Player', 'name', 'alpha'],
        ['2', 'Kills', 'killsOnFoot', '★ desc'],
        ['3', 'Veh Kills', 'vehKillsFoot', 'desc'],
        ['4', 'Deaths', 'deathsOnFoot', 'desc'],
        ['5', 'K/D', 'kdFoot', 'desc (green ≥2, red <0.8)'],
        ['6', 'TK', 'tkOnFoot', 'desc (red if >0)'],
        ['7', 'Suicides', 'suicides', 'desc'],
        ['8', 'Shots', 'shotsOnFoot', 'desc'],
        ['9', 'Hits Taken', 'hitsOnFoot', 'desc'],
        ['10', 'Shots/Kill', 'spkFoot', 'asc'],
        ['11', 'Avg Dist (m)', 'avgDistFoot', 'desc'],
        ['12', 'Longest (m)', 'maxLongestFoot', 'desc'],
        ['13', 'Missions', 'missionCount', 'desc'],
    ]
)
doc.add_paragraph()
add_para(doc,
    'Rows with tkOnFoot > 0 receive the .tk-row CSS class (light pink background). '
    'Rank medals (🥇🥈🥉) are assigned by killsOnFoot position. '
    'Clicking a row opens the player detail modal. '
    'The 📊 icon in the Player column opens the full career page.',
    space_after=6)

add_heading(doc, '6.5 renderVehicleTable()', 2)
add_para(doc,
    'Renders the Vehicle Combat Stats table. Only players with at least one vehicle '
    'kill or death are included. Column structure mirrors the infantry table but for '
    'vehicle metrics.',
    space_after=6)

add_heading(doc, '6.6 renderTableHead()', 2)
add_para(doc,
    'Renders clickable <th> elements with sort state indicators. '
    'The active sort column shows ▲ (ascending) or ▼ (descending); '
    'inactive columns show ⇅. Clicks call _sortInf() or _sortVeh() which toggle '
    'direction if the same column is clicked twice.',
    space_after=6)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  7. PLAYER DETAIL VIEWS
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '7. Player Detail Views', 1)

add_heading(doc, '7.1 Player Detail Modal', 2)
add_para(doc,
    'Triggered by clicking any row in either table. Opens a floating overlay '
    '(.modal-overlay.open) showing a full career breakdown for the selected player.',
    space_after=4)
add_table(doc,
    ['Element', 'Content'],
    [
        ['Header', 'Player name, mission count, active date range, top role'],
        ['Maximize button (⛶)', 'Closes modal, opens full career page for same player'],
        ['Close button (✕)', 'Removes .open class; also closed by pressing ESC or clicking overlay'],
        ['Body', 'buildCareerStatsHTML() output — 4 sections (see §7.3)'],
    ]
)

doc.add_paragraph()
add_heading(doc, '7.2 Full Career Page', 2)
add_para(doc,
    'Triggered by: right-clicking a player pill, clicking the 📊 icon on a table row, '
    'or clicking the maximize button in the modal.',
    space_after=4)
add_para(doc, 'When the career page opens:', space_after=2)
bullet(doc, 'selectedPlayers is set to a one-element Set for the player.')
bullet(doc, 'applyFilters() is called — all tables and stats re-filter to this player.')
bullet(doc, 'The stats bar, awards row, chart, and filter panel are hidden.')
bullet(doc, 'The career header and career stats panel are shown.')
bullet(doc, 'The page scrolls to the top.')
doc.add_paragraph()
add_para(doc, 'Career page header shows:', space_after=2)
bullet(doc, 'Player name (large, Oswald font)')
bullet(doc, 'Combat Missions count  ·  Active date range (DD/MM/YYYY – DD/MM/YYYY)  ·  Top Role (count)')
doc.add_paragraph()
add_para(doc,
    'Clicking "← Back to Leaderboard" calls closeCareerPage(), which resets selectedPlayers to null, '
    're-shows all leaderboard elements, and calls applyFilters().',
    space_after=6)

add_heading(doc, '7.3 buildCareerStatsHTML() — Content Sections', 2)
add_table(doc,
    ['Section', 'Content'],
    [
        ['1 — Overall Infantry Stats',
         'Grid of stat cards: Kills, Deaths, K/D, TK, Shots, Hits Taken, Shots/Kill, Avg Dist, Longest Kill, Missions, Suicides'],
        ['2 — Weapon Kill Breakdown',
         'Top 15 weapons sorted by kill count; proportional bar chart with kill count and percentage of total kills'],
        ['3 — Best Single Mission',
         'Mission with the highest kill count; shows kills, K/D, TK, and longest kill for that mission'],
        ['4 — Kill Breakdown by Mission',
         'Table of all missions sorted by kills desc; columns: Mission, Date, Kills, Deaths, K/D, TK, Longest Kill (m)'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  8. COMPLETE FUNCTION REFERENCE
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '8. Complete Function Reference', 1)

functions = [
    ('NUM(v)', 'Utility', 'v: any', 'number', 'Safe parseFloat with 0 fallback.'),
    ('normalizeRole(role)', 'Utility', 'role: string', 'string', 'Strips @username, squad letters, colour/phonetic suffixes from role names.'),
    ('missionDate(name)', 'Utility', 'name: string', 'string', 'Extracts (YYYY-MM-DD) from mission name and reformats to DD/MM/YYYY.'),
    ('isJointOp(filename)', 'Utility', 'filename: string', 'boolean', 'Returns true if filename date is the last Saturday or its following Sunday of the month.'),
    ('parseCSVLine(line)', 'Parsing', 'line: string', 'string[]', 'RFC-4180 CSV parser that handles quoted fields with embedded commas.'),
    ('buildAggregates()', 'Aggregation', '—', 'void', 'Loops rawRows and builds the aggPlayers dictionary of career totals. Called once at load.'),
    ('buildUI()', 'Lifecycle', '—', 'void', 'Shows content div, hides loader, calls buildFilters() and applyFilters().'),
    ('buildFilters()', 'Filters', '—', 'void', 'Wires all filter UI event listeners (event type, reset, search inputs, Zeus dropdown).'),
    ('getEventFilteredRows()', 'Filters', '—', 'Object[]', 'Returns rawRows filtered by showJointOps / showRegularEvents.'),
    ('refreshPills()', 'Filters', '—', 'void', 'Recomputes available missions and players, then re-renders both pill containers.'),
    ('renderCareerPills(cId, items, sId)', 'Filters', 'containerId, items[], searchId', 'void', 'Renders player pills. Left-click = filter; right-click = career page.'),
    ('renderSelectablePills(cId, items, sel, sId, cb)', 'Filters', 'containerId, items[], selectedSet, searchId, onSelect()', 'void', 'Generic multi-select pill renderer. null selectedSet means "all active".'),
    ('filterToPlayer(name)', 'Filters', 'name: string', 'void', 'Toggles player in/out of selectedPlayers; null if set becomes empty.'),
    ('filterChanged()', 'Filters', '—', 'void', 'Calls applyFilters(). Entry point for all filter state changes.'),
    ('applyFilters()', 'Filters', '—', 'void', 'Core pipeline: filter rawRows → re-aggregate → Zeus filter → render all.'),
    ('renderStats()', 'Rendering', '—', 'void', 'Renders aggregate stat cards to #statsBar.'),
    ('renderChart()', 'Rendering', '—', 'void', 'Renders top-10 infantry killers horizontal bar chart to #chartBars.'),
    ('renderLeader()', 'Rendering', '—', 'void', 'Renders the 4 award cards (Executioner, Pro Sniper, Perfect Aim, K/D Player).'),
    ('renderInfantryTable()', 'Rendering', '—', 'void', 'Sorts filteredPlayers and writes infantry table rows/header to DOM.'),
    ('renderVehicleTable()', 'Rendering', '—', 'void', 'Sorts vehicle-active players and writes vehicle table rows/header to DOM.'),
    ('renderTableHead(hId, cols, sCol, sAsc, tId)', 'Rendering', 'headId, cols[], sortCol, sortAsc, tableId', 'void', 'Renders <th> elements with ▲▼⇅ sort indicators.'),
    ('sortPlayers(arr, col, asc, cols)', 'Rendering', 'arr, colIdx, ascending, cols[]', 'Object[]', 'Generic sort; uses cols[col].sortKey; nulls sorted to end.'),
    ('openPlayerModal(name)', 'Modal', 'name: string', 'void', 'Opens the detail modal overlay for the given player.'),
    ('closeModal(e)', 'Modal', 'event', 'void', 'Closes modal on overlay click or ESC keypress.'),
    ('openCareerPage(name)', 'Career', 'name: string', 'void', 'Switches to full-screen career view; sets selectedPlayers and re-filters.'),
    ('closeCareerPage()', 'Career', '—', 'void', 'Restores leaderboard view; clears selectedPlayers and re-filters.'),
    ('buildCareerStatsHTML(p)', 'Career', 'p: PlayerObject', 'string', 'Generates HTML string for 4-section career stats page/modal body.'),
    ('mStat(val, lbl, sub)', 'Helpers', 'val, label, subtitle?', 'string', 'Returns HTML for a single stat card element.'),
    ('kdClass(v)', 'Helpers', 'v: number', 'string', '"kd-good" if ≥2, "kd-bad" if <0.8, else "".'),
    ('tkClass(v)', 'Helpers', 'v: number', 'string', '"tk-cell" if >0, else "".'),
    ('window._sortInf(col)', 'Sort', 'col: number', 'void', 'Infantry table header click handler; toggles direction if same column.'),
    ('window._sortVeh(col)', 'Sort', 'col: number', 'void', 'Vehicle table header click handler; toggles direction if same column.'),
]

add_table(doc,
    ['Function', 'Category', 'Parameters', 'Returns', 'Description'],
    functions
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  9. CSS DESIGN SYSTEM
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '9. CSS Design System', 1)

add_heading(doc, '9.1 CSS Custom Properties (Variables)', 2)
add_table(doc,
    ['Variable', 'Value', 'Usage'],
    [
        ['--charcoal', '#424242', 'Primary dark colour — headers, section titles, table headers'],
        ['--charcoal-light', '#5a5a5a', 'Hover states, secondary elements'],
        ['--bg', '#f4f4f4', 'Page background'],
        ['--white', '#ffffff', 'Card and panel backgrounds'],
        ['--green', '#388e3c', 'Good K/D (≥ 2), positive indicators'],
        ['--red', '#c62828', 'Teamkills, bad K/D (< 0.8), error states'],
        ['--gold', '#f9a825', 'Rank #1 medal, top stat highlights'],
        ['--silver', '#9e9e9e', 'Rank #2 medal'],
        ['--bronze', '#795548', 'Rank #3 medal'],
        ['--border', '#e0e0e0', 'Card/panel borders, table cell borders'],
        ['--row-hover', '#f0f0f0', 'Table row hover background'],
        ['--tk-bg', '#fff3f3', 'Teamkill row background (light pink)'],
    ]
)

doc.add_paragraph()
add_heading(doc, '9.2 Typography', 2)
add_table(doc,
    ['Font Family', 'Weights', 'Usage'],
    [
        ['Oswald', '400, 600, 700', 'Section headings, award card titles, large stat values, player name on career page'],
        ['Open Sans', '400, 600', 'Body text, table cells, filter labels, pill text'],
        ['Courier New', '400', 'Code blocks (documentation only)'],
    ]
)

doc.add_paragraph()
add_heading(doc, '9.3 Key CSS Classes', 2)
add_table(doc,
    ['Class', 'Element', 'Purpose'],
    [
        ['.stat-card', '<div>', 'Top stats bar card; flex-grow 1, min-width 110px'],
        ['.filter-panel', '<div>', 'White rounded panel containing all filter controls'],
        ['.pill', '<span>', 'Filter pill; border-radius 20px; cursor pointer'],
        ['.pill.active', '<span>', 'Dark background (#424242) / white text — active selection'],
        ['.award-card', '<div>', 'Award card base; rounded; flex column'],
        ['.ac-kills / .ac-longest / .ac-spk / .ac-kda', '<div>', 'Award card colour variants (background differs per award)'],
        ['.kd-good', '<td>', 'Green text for K/D ≥ 2'],
        ['.kd-bad', '<td>', 'Red text for K/D < 0.8'],
        ['.tk-cell', '<td>', 'Red text for teamkill count > 0'],
        ['.tk-row', '<tr>', 'Light pink (#fff3f3) row background for players with TKs'],
        ['.rank-gold / .rank-silver / .rank-bronze', '<td>', 'Medal emoji colouring for top-3 ranks by kills'],
        ['.modal-overlay', '<div>', 'Full-screen semi-transparent overlay; .open class shows it'],
        ['.career-header', '<div>', 'Career page top bar with back button and player name'],
        ['.career-icon-btn', '<button>', 'The 📊 icon in player column that opens career page'],
        ['.modal-maximize', '<button>', 'The ⛶ button in modal header to expand to career page'],
        ['.weapon-track / .weapon-fill', '<div>', 'Weapon breakdown bar chart; fill width set via inline style percentage'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  10. DOM STRUCTURE
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '10. DOM Structure', 1)

add_heading(doc, '10.1 Top-level Elements', 2)
add_table(doc,
    ['ID / Selector', 'Tag', 'Purpose'],
    [
        ['#loading', 'div', 'Shown during CSV fetch; hidden on success'],
        ['#error', 'div', 'Shown on fetch failure with error message'],
        ['#content', 'div', 'Main app container; hidden until data loads'],
        ['#statsBar', 'div', 'Row of aggregate stat cards'],
        ['#awardsRow', 'div', 'Row of 4 award cards'],
        ['#chartWrap', 'div', 'Top-10 killers bar chart wrapper'],
        ['#chartBars', 'div', 'Bar chart inner container'],
        ['.filter-panel', 'div', 'Filter controls (event type, search, pills, zeus)'],
        ['#playerSearch', 'input', 'Player name search box'],
        ['#missionSearch', 'input', 'Mission name search box'],
        ['#playerPills', 'div', 'Player pill container'],
        ['#missionPills', 'div', 'Mission pill container'],
        ['#filterCount', 'span', 'Active filter summary text'],
        ['#eventTypeSelect', 'select', 'Event type dropdown (Both / Joint Op / Regular Op)'],
        ['#zeusFilterSelect', 'select', 'Zeus filter dropdown (All / No Zeus / Zeus Only)'],
        ['#resetFilters', 'button', 'Resets all filters to defaults'],
        ['#careerHeader', 'div', 'Career page header (hidden in leaderboard mode)'],
        ['#careerPlayerName', 'div', 'Player name in career header'],
        ['#careerPlayerSub', 'div', 'Missions · Active dates · Top role subtitle'],
        ['#careerStats', 'div', 'Career page body content from buildCareerStatsHTML()'],
        ['#infantryHead', 'thead', 'Infantry table header row'],
        ['#infantryBody', 'tbody', 'Infantry table data rows'],
        ['#vehicleHead', 'thead', 'Vehicle table header row'],
        ['#vehicleBody', 'tbody', 'Vehicle table data rows'],
        ['#playerModal', 'div.modal-overlay', 'Player detail modal overlay'],
        ['#modalPlayerName', 'h2', 'Player name in modal header'],
        ['#modalPlayerSub', 'div', 'Mission count / date range / role in modal header'],
        ['#modalMaximizeBtn', 'button', 'Expand modal to full career page'],
        ['#modalBody', 'div', 'Modal content (career HTML from buildCareerStatsHTML)'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  11. DEPLOYMENT
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '11. Deployment', 1)

add_heading(doc, '11.1 GitHub Actions Pipeline', 2)
add_para(doc,
    'A GitHub Actions workflow deploys the application automatically on every push.',
    space_after=4)
add_table(doc,
    ['Branch Pattern', 'Deployed URL', 'Purpose'],
    [
        ['main', 'https://loltorres9.github.io/TFP_KillTracker/', 'Production'],
        ['claude/*', 'https://loltorres9.github.io/TFP_KillTracker/dev/', 'Preview / development'],
    ]
)

doc.add_paragraph()
add_heading(doc, '11.2 GitHub Pages Configuration', 2)
add_para(doc,
    'GitHub Pages must be configured to serve from the gh-pages branch. '
    'Navigate to repository Settings → Pages → Source and select the gh-pages branch.',
    space_after=6)

add_heading(doc, '11.3 Publishing a Preview to Production', 2)
numbered(doc, 'Develop on the claude/* branch.')
numbered(doc, 'Open a pull request from the claude/* branch into main.')
numbered(doc, 'Review and merge the pull request.')
numbered(doc, 'The GitHub Actions workflow triggers and deploys to production.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  12. EXTENDING THE APPLICATION
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '12. Extending the Application', 1)

add_heading(doc, '12.1 Adding a New Stat Column', 2)
numbered(doc, 'Add the raw column to the Google Sheet and publish the updated CSV.')
numbered(doc, 'In buildAggregates(), add an initialiser (e.g. p.newStat = 0) and an accumulator.')
numbered(doc, 'In applyFilters(), mirror the accumulation in the tempAgg block.')
numbered(doc, 'Add a column definition object to INF_COLS or VEH_COLS: { label, key, numeric, sortKey, render }.')
numbered(doc, 'The table renderers will pick up the new column automatically.')

add_heading(doc, '12.2 Adding a New Award Card', 2)
numbered(doc, 'Add an HTML award-card div in the #awardsRow section of index.html.')
numbered(doc, 'Give it a unique ID for each data element (name, stat, sub-stat).')
numbered(doc, 'In renderLeader(), sort/filter filteredPlayers to find the winner.')
numbered(doc, 'Write the winner\'s stats to the new DOM elements.')

add_heading(doc, '12.3 Pointing to a Different Google Sheet', 2)
numbered(doc, 'Publish the new sheet: File → Share → Publish to web → Select sheet → CSV → Publish.')
numbered(doc, 'Copy the published URL.')
numbered(doc, 'In index.html, update the CSV_URL constant at the top of the <script> block.')

add_heading(doc, '12.4 Adding Persistence (localStorage)', 2)
add_para(doc,
    'The application currently holds no persistent state. To save filter preferences across page loads:',
    space_after=4)
bullet(doc, 'Write selected state to localStorage in filterChanged().')
bullet(doc, 'Read and restore state in buildFilters() before calling refreshPills().')
bullet(doc, 'Use JSON.stringify/parse for Sets.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  APPENDIX
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'Appendix A — Initialization Sequence', 1)
add_para(doc, 'The following sequence occurs on every page load:', space_after=4)

steps = [
    ('Page Load', 'Browser requests index.html; CSS and JS parse inline.'),
    ('DOMContentLoaded', 'JavaScript execution begins; loading spinner displayed.'),
    ('fetch(CSV_URL)', 'HTTP GET to Google Sheets export URL.'),
    ('parseCSVLine() × N', 'Each row parsed; rawRows[] populated.'),
    ('buildAggregates()', 'aggPlayers{} populated with career totals.'),
    ('buildUI()', 'Calls buildFilters() and applyFilters().'),
    ('buildFilters()', 'Event listeners attached; refreshPills() called.'),
    ('applyFilters()', 'filteredPlayers[] built; five render functions called.'),
    ('renderStats()', '#statsBar populated.'),
    ('renderChart()', '#chartBars populated.'),
    ('renderLeader()', '#awardsRow populated.'),
    ('renderInfantryTable()', '#infantryBody + #infantryHead populated.'),
    ('renderVehicleTable()', '#vehicleBody + #vehicleHead populated.'),
    ('UI Ready', 'Loading spinner hidden; content shown.'),
]

add_table(doc,
    ['Step', 'Function / Event', 'Description'],
    [[str(i+1), fn, desc] for i, (fn, desc) in enumerate(steps)]
)

doc.add_paragraph()
add_heading(doc, 'Appendix B — Keyboard Shortcuts', 1)
add_table(doc,
    ['Key', 'Context', 'Action'],
    [
        ['ESC', 'Modal open', 'Close the player detail modal'],
        ['Click overlay', 'Modal open', 'Close the player detail modal'],
    ]
)

doc.add_paragraph()
add_heading(doc, 'Appendix C — Browser Compatibility', 1)
add_para(doc,
    'The application uses standard ES6+ features: arrow functions, template literals, '
    'destructuring, Set, Map, fetch(), and async/await. It is compatible with all modern '
    'browsers (Chrome 80+, Firefox 75+, Safari 13+, Edge 80+). '
    'Internet Explorer is not supported.',
    space_after=6)

# ── Save ──────────────────────────────────────────────────────────────────────

out_path = '/home/user/TFP_KillTracker/TFP_KillTracker_Technical_Documentation.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
